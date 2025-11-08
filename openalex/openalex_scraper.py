import requests
import json
from typing import Dict, List, Optional, Union
import time


class OpenAlexScraper:
    
    
    BASE_URL = "https://api.openalex.org"
    ENTITY_TYPES = ['works', 'authors', 'sources', 'institutions', 'topics', 'publishers', 'funders']
    
    def __init__(self, email: Optional[str] = None, api_key: Optional[str] = None, rate_limit_delay: float = 0.1):
        
        self.email = email
        self.api_key = api_key
        self.rate_limit_delay = rate_limit_delay
        self.session = requests.Session()
        
        # Set User-Agent header with email (official best practice)
        if email:
            self.session.headers.update({
                "User-Agent": f"mailto:{email}"
            })
        else:
            self.session.headers.update({
                "User-Agent": "OpenAlexScraper/1.0"
            })
    
    def _make_request(self, endpoint: str, params: Optional[Dict] = None) -> Dict:
        """
        Make a request to the OpenAlex API with error handling
        
        Args:
            endpoint: API endpoint (e.g., '/works', '/authors')
            params: Query parameters
            
        Returns:
            JSON response as dictionary, empty dict on error
        """
        url = f"{self.BASE_URL}{endpoint}"
        try:
            response = self.session.get(url, params=params, timeout=30)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.Timeout:
            print(f"Error: Request timeout for {endpoint}")
            return {}
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 429:
                print(f"Error: Rate limit exceeded. Waiting 60 seconds...")
                time.sleep(60)
                return self._make_request(endpoint, params)
            print(f"Error: HTTP {e.response.status_code} for {endpoint}")
            return {}
        except requests.exceptions.RequestException as e:
            print(f"Error making request to {endpoint}: {e}")
            return {}
    
    def _normalize_id(self, entity_id: str) -> str:
        """Extract ID from full URL or return as-is"""
        if entity_id.startswith('http'):
            return entity_id.split('/')[-1]
        return entity_id
    
    # ==================== Generic Entity Methods ====================
    
    def get_entity(self, entity_type: str, entity_id: str) -> Dict:
        """
        Get a single entity by type and ID
        
        Args:
            entity_type: Entity type (works, authors, sources, institutions, topics, publishers, funders)
            entity_id: OpenAlex entity ID (e.g., 'W2741803907' or full URL)
            
        Returns:
            Entity data as dictionary
        """
        if entity_type not in self.ENTITY_TYPES:
            raise ValueError(f"Invalid entity type. Must be one of: {self.ENTITY_TYPES}")
        
        entity_id = self._normalize_id(entity_id)
        params = {}
        if self.api_key:
            params['api_key'] = self.api_key
        return self._make_request(f"/{entity_type}/{entity_id}", params=params if params else None)
    
    def search_entities(self,
                       entity_type: str,
                     query: Optional[str] = None,
                     filter_params: Optional[Dict] = None,
                       sort_params: Optional[Dict] = None,
                     per_page: int = 25,
                       page: Optional[int] = None,
                       cursor: Optional[str] = None) -> Dict:
        """
        Search for entities with filtering, sorting, and pagination
        
        Args:
            entity_type: Entity type (works, authors, sources, institutions, topics, publishers, funders)
            query: Search query string
            filter_params: Filter parameters (e.g., {'author.id': 'A2208157607'})
            sort_params: Sort parameters (e.g., {'cited_by_count': 'desc'})
            per_page: Number of results per page (max 200)
            page: Page number (for page-based pagination)
            cursor: Cursor string (for cursor-based pagination, more efficient)
            
        Returns:
            Search results as dictionary
        """
        if entity_type not in self.ENTITY_TYPES:
            raise ValueError(f"Invalid entity type. Must be one of: {self.ENTITY_TYPES}")
        
        params = {
            'per_page': min(per_page, 200)
        }
        
        # Add API key if provided (for Premium plan)
        if self.api_key:
            params['api_key'] = self.api_key
        
        if query:
            params['search'] = query
        
        if filter_params:
            # OpenAlex uses filter=key:value format (colon separator)
            # Multiple filters can be combined with commas: filter=key1:value1,key2:value2
            filter_strings = [f'{key}:{value}' for key, value in filter_params.items()]
            params['filter'] = ','.join(filter_strings)
        
        if sort_params:
            for key, value in sort_params.items():
                params[f'sort.{key}'] = value
        
        if cursor:
            params['cursor'] = cursor
        elif page:
            params['page'] = page
        
        time.sleep(self.rate_limit_delay)
        return self._make_request(f"/{entity_type}", params=params)
    
    def get_all_entities(self,
                        entity_type: str,
                        query: Optional[str] = None,
                        filter_params: Optional[Dict] = None,
                        sort_params: Optional[Dict] = None,
                        max_results: Optional[int] = None,
                        use_cursor: bool = True) -> List[Dict]:
        """
        Get all entities matching criteria using cursor-based pagination
        
        Args:
            entity_type: Entity type (works, authors, sources, institutions, topics, publishers, funders)
            query: Search query string
            filter_params: Filter parameters
            sort_params: Sort parameters
            max_results: Maximum number of results to fetch (None for all)
            use_cursor: Use cursor-based pagination (more efficient) or page-based
            
        Returns:
            List of entity dictionaries
        """
        all_entities = []
        cursor = None
        page = 1
        per_page = 200  # Use max per_page for efficiency
        
        while True:
            if use_cursor:
                results = self.search_entities(
                    entity_type=entity_type,
                    query=query,
                    filter_params=filter_params,
                    sort_params=sort_params,
                    per_page=per_page,
                    cursor=cursor
                )
            else:
                results = self.search_entities(
                    entity_type=entity_type,
                    query=query,
                    filter_params=filter_params,
                    sort_params=sort_params,
                    per_page=per_page,
                    page=page
                )
            
            if 'results' not in results or not results['results']:
                break
            
            all_entities.extend(results['results'])
            
            # Check if we've reached max_results
            if max_results and len(all_entities) >= max_results:
                all_entities = all_entities[:max_results]
                break
            
            # Get next cursor or page
            if use_cursor:
                cursor = results.get('meta', {}).get('next_cursor')
                if not cursor:
                    break
            else:
                if len(results['results']) < per_page:
                    break
                page += 1
        
        return all_entities
    
    # ==================== Works Methods ====================
    
    def get_work(self, work_id: str) -> Dict:
        """Get a single work by ID"""
        return self.get_entity('works', work_id)
    
    def search_works(self, 
                     query: Optional[str] = None,
                     filter_params: Optional[Dict] = None,
                     sort_params: Optional[Dict] = None,
                     per_page: int = 25,
                     page: Optional[int] = None,
                     cursor: Optional[str] = None) -> Dict:
        """Search for works"""
        return self.search_entities('works', query, filter_params, sort_params, per_page, page, cursor)
    
    def get_all_works(self,
                     query: Optional[str] = None,
                     filter_params: Optional[Dict] = None,
                     sort_params: Optional[Dict] = None,
                     max_results: Optional[int] = None,
                     use_cursor: bool = True) -> List[Dict]:
        """Get all works matching criteria"""
        return self.get_all_entities('works', query, filter_params, sort_params, max_results, use_cursor)
    
    # ==================== Authors Methods ====================
    
    def get_author(self, author_id: str) -> Dict:
        """Get a single author by ID"""
        return self.get_entity('authors', author_id)
    
    def search_authors(self,
                      query: Optional[str] = None,
                      filter_params: Optional[Dict] = None,
                      sort_params: Optional[Dict] = None,
                      per_page: int = 25,
                      page: Optional[int] = None,
                      cursor: Optional[str] = None) -> Dict:
        """Search for authors"""
        return self.search_entities('authors', query, filter_params, sort_params, per_page, page, cursor)
    
    def get_all_authors(self,
                       query: Optional[str] = None,
                       filter_params: Optional[Dict] = None,
                       sort_params: Optional[Dict] = None,
                       max_results: Optional[int] = None,
                       use_cursor: bool = True) -> List[Dict]:
        """Get all authors matching criteria"""
        return self.get_all_entities('authors', query, filter_params, sort_params, max_results, use_cursor)
    
    # ==================== Institutions Methods ====================
    
    def get_institution(self, institution_id: str) -> Dict:
        """Get a single institution by ID"""
        return self.get_entity('institutions', institution_id)
    
    def search_institutions(self,
                           query: Optional[str] = None,
                           filter_params: Optional[Dict] = None,
                           sort_params: Optional[Dict] = None,
                           per_page: int = 25,
                           page: Optional[int] = None,
                           cursor: Optional[str] = None) -> Dict:
        """Search for institutions"""
        return self.search_entities('institutions', query, filter_params, sort_params, per_page, page, cursor)
    
    def get_all_institutions(self,
                            query: Optional[str] = None,
                            filter_params: Optional[Dict] = None,
                            sort_params: Optional[Dict] = None,
                            max_results: Optional[int] = None,
                            use_cursor: bool = True) -> List[Dict]:
        """Get all institutions matching criteria"""
        return self.get_all_entities('institutions', query, filter_params, sort_params, max_results, use_cursor)
    
    # ==================== Sources Methods ====================
    
    def get_source(self, source_id: str) -> Dict:
        """Get a single source (venue/journal) by ID"""
        return self.get_entity('sources', source_id)
    
    def search_sources(self,
                      query: Optional[str] = None,
                      filter_params: Optional[Dict] = None,
                      sort_params: Optional[Dict] = None,
                      per_page: int = 25,
                      page: Optional[int] = None,
                      cursor: Optional[str] = None) -> Dict:
        """Search for sources (venues/journals)"""
        return self.search_entities('sources', query, filter_params, sort_params, per_page, page, cursor)
    
    # ==================== Topics Methods ====================
    
    def get_topic(self, topic_id: str) -> Dict:
        """Get a single topic by ID"""
        return self.get_entity('topics', topic_id)
    
    def search_topics(self,
                     query: Optional[str] = None,
                     filter_params: Optional[Dict] = None,
                     sort_params: Optional[Dict] = None,
                     per_page: int = 25,
                     page: Optional[int] = None,
                     cursor: Optional[str] = None) -> Dict:
        """Search for topics"""
        return self.search_entities('topics', query, filter_params, sort_params, per_page, page, cursor)
    
    # ==================== Publishers Methods ====================
    
    def get_publisher(self, publisher_id: str) -> Dict:
        """Get a single publisher by ID"""
        return self.get_entity('publishers', publisher_id)
    
    def search_publishers(self,
                         query: Optional[str] = None,
                         filter_params: Optional[Dict] = None,
                         sort_params: Optional[Dict] = None,
                         per_page: int = 25,
                         page: Optional[int] = None,
                         cursor: Optional[str] = None) -> Dict:
        """Search for publishers"""
        return self.search_entities('publishers', query, filter_params, sort_params, per_page, page, cursor)
    
    # ==================== Funders Methods ====================
    
    def get_funder(self, funder_id: str) -> Dict:
        """Get a single funder by ID"""
        return self.get_entity('funders', funder_id)
    
    def search_funders(self,
                      query: Optional[str] = None,
                      filter_params: Optional[Dict] = None,
                      sort_params: Optional[Dict] = None,
                      per_page: int = 25,
                      page: Optional[int] = None,
                      cursor: Optional[str] = None) -> Dict:
        """Search for funders"""
        return self.search_entities('funders', query, filter_params, sort_params, per_page, page, cursor)
    
    # ==================== Utility Methods ====================
    
    def scrape_works_by_topic(self, topic: str, max_results: int = 100) -> List[Dict]:
        """
        Scrape works related to a specific topic (deprecated - use get_all_works instead)
        
        Args:
            topic: Topic name to search for
            max_results: Maximum number of results to fetch
            
        Returns:
            List of work dictionaries
        """
        return self.get_all_works(query=topic, max_results=max_results)
    
    def get_works_by_author(self, author_id: str, max_results: Optional[int] = None) -> List[Dict]:
        """
        Get all works by a specific author
        
        Args:
            author_id: OpenAlex author ID
            max_results: Maximum number of results to fetch
            
        Returns:
            List of work dictionaries
        """
        author_id = self._normalize_id(author_id)
        filter_params = {'author.id': author_id}
        return self.get_all_works(filter_params=filter_params, max_results=max_results)
    
    def get_works_by_institution(self, institution_id: str, max_results: Optional[int] = None) -> List[Dict]:
        """
        Get all works from a specific institution
        
        Args:
            institution_id: OpenAlex institution ID
            max_results: Maximum number of results to fetch
            
        Returns:
            List of work dictionaries
        """
        institution_id = self._normalize_id(institution_id)
        filter_params = {'institutions.id': institution_id}
        return self.get_all_works(filter_params=filter_params, max_results=max_results)
    
    def get_works_by_year(self, year: int, max_results: Optional[int] = None) -> List[Dict]:
        """
        Get all works published in a specific year
        
        Args:
            year: Publication year
            max_results: Maximum number of results to fetch
            
        Returns:
            List of work dictionaries
        """
        filter_params = {'publication_year': year}
        return self.get_all_works(filter_params=filter_params, max_results=max_results)
    
    def save_to_json(self, data: Union[List[Dict], Dict], filename: str):
        """
        Save scraped data to a JSON file
        
        Args:
            data: List of dictionaries or single dictionary to save
            filename: Output filename
        """
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        if isinstance(data, list):
            print(f"Saved {len(data)} records to {filename}")
        else:
            print(f"Saved data to {filename}")
    
    def save_to_csv(self, data: List[Dict], filename: str, flatten_nested: bool = True):
        """
        Save scraped data to a CSV file
        
        Args:
            data: List of dictionaries to save
            filename: Output filename
            flatten_nested: Whether to flatten nested dictionaries (basic implementation)
        """
        try:
            import csv
        except ImportError:
            print("CSV module not available. Install Python standard library.")
            return
        
        if not data:
            print("No data to save")
            return
        
        # Simple CSV writer (for complex nested data, consider pandas)
        fieldnames = set()
        for item in data:
            fieldnames.update(item.keys())
        
        fieldnames = sorted(list(fieldnames))
        
        with open(filename, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction='ignore')
            writer.writeheader()
            
            for item in data:
                # Convert nested dicts/lists to strings for CSV
                row = {}
                for key, value in item.items():
                    if isinstance(value, (dict, list)):
                        row[key] = json.dumps(value) if flatten_nested else str(value)
                    else:
                        row[key] = value
                writer.writerow(row)
        
        print(f"Saved {len(data)} records to {filename}")
    
    def save_to_excel(self, data: List[Dict], filename: str):
        """
        Save scraped data to an Excel file (.xlsx)
        
        Args:
            data: List of dictionaries to save
            filename: Output filename (should end with .xlsx)
        """
        try:
            import pandas as pd
        except ImportError:
            print("Error: pandas and openpyxl required for Excel export.")
            print("Install with: pip install pandas openpyxl")
            return
        
        if not data:
            print("No data to save")
            return
        
        try:
            # Convert to DataFrame
            df = pd.json_normalize(data)
            df.to_excel(filename, index=False, engine='openpyxl')
            print(f"Saved {len(data)} records to {filename}")
        except Exception as e:
            print(f"Error saving to Excel: {e}")
    
    def save_to_word(self, data: List[Dict], filename: str):
        """
        Save scraped data to a Word document (.docx)
        
        Args:
            data: List of dictionaries to save
            filename: Output filename (should end with .docx)
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            print("Error: python-docx required for Word export.")
            print("Install with: pip install python-docx")
            return
        
        if not data:
            print("No data to save")
            return
        
        try:
            doc = Document()
            doc.add_heading('OpenAlex Scraped Data', 0)
            doc.add_paragraph(f'Total Records: {len(data)}')
            doc.add_paragraph('')
            
            for idx, item in enumerate(data, 1):
                doc.add_heading(f'Record {idx}', level=1)
                
                # Add key-value pairs
                for key, value in item.items():
                    if value is None:
                        continue
                    
                    # Format the value
                    if isinstance(value, (dict, list)):
                        value_str = json.dumps(value, indent=2, ensure_ascii=False)
                    else:
                        value_str = str(value)
                    
                    # Add to document
                    p = doc.add_paragraph()
                    p.add_run(f'{key}: ').bold = True
                    p.add_run(value_str[:5000])  # Limit length
                
                doc.add_paragraph('')  # Add spacing
            
            doc.save(filename)
            print(f"Saved {len(data)} records to {filename}")
        except Exception as e:
            print(f"Error saving to Word: {e}")
    
    def _extract_pdf_from_html(self, html_content: str, base_url: str) -> Optional[str]:
        """
        Extract PDF URL from HTML content
        
        Args:
            html_content: HTML content of the page
            base_url: Base URL for resolving relative links
            
        Returns:
            PDF URL if found, None otherwise
        """
        try:
            import re
            from urllib.parse import urljoin, urlparse
            
            # Common PDF link patterns in HTML
            pdf_patterns = [
                r'href=["\']([^"\']*\.pdf[^"\']*)["\']',
                r'href=["\']([^"\']*pdf[^"\']*)["\']',
                r'["\']([^"\']*download[^"\']*\.pdf[^"\']*)["\']',
                r'["\']([^"\']*\.pdf[^"\']*download[^"\']*)["\']',
            ]
            
            for pattern in pdf_patterns:
                matches = re.findall(pattern, html_content, re.IGNORECASE)
                for match in matches:
                    if '.pdf' in match.lower() or 'pdf' in match.lower():
                        # Resolve relative URLs
                        pdf_url = urljoin(base_url, match)
                        return pdf_url
            
            return None
        except Exception:
            return None
    
    def download_pdf(self, pdf_url: str, filename: str, verbose: bool = False) -> bool:
        
        if not pdf_url:
            return False
        
        # Try direct URL first
        if self._try_download_pdf(pdf_url, filename, verbose):
            return True
        
        # If it's a DOI, try DOI-based PDF patterns
        if 'doi.org' in pdf_url or pdf_url.startswith('10.'):
            doi_clean = pdf_url.replace('https://doi.org/', '').replace('http://dx.doi.org/', '').replace('https://', '').replace('http://', '')
            # Try DOI resolver patterns
            doi_patterns = [
                f"https://doi.org/{doi_clean}.pdf",
                f"https://www.doi.org/{doi_clean}.pdf",
                f"https://dx.doi.org/{doi_clean}.pdf",
            ]
            for pattern in doi_patterns:
                if self._try_download_pdf(pattern, filename, verbose):
                    return True
        
        # If it's a landing page or article page, try common PDF patterns
        if any(x in pdf_url.lower() for x in ['/article/', '/view/', 'sciencedirect', 'springer', 'ieee', 'elsevier', 'tandfonline', 'wiley', 'acm', 'arxiv', 'biorxiv', 'medrxiv', 'researchgate', 'academia']):
            # Try many more PDF patterns
            pdf_patterns = []
            
            # Common patterns
            if '/article/' in pdf_url:
                pdf_patterns.extend([
                    pdf_url.replace('/article/', '/article/pdf/'),
                    pdf_url.replace('/article/', '/pdf/'),
                    pdf_url.replace('/article/', '/article/epdf/'),
                    pdf_url.replace('/article/', '/article/fulltext/') + '/pdf',
                ])
            if '/view/' in pdf_url:
                pdf_patterns.extend([
                    pdf_url.replace('/view/', '/download/'),
                    pdf_url.replace('/view/', '/download/') + '/pdf',
                ])
            if '/chapter/' in pdf_url:
                pdf_patterns.extend([
                    pdf_url.replace('/chapter/', '/chapter/pdf/'),
                ])
            
            # Publisher-specific patterns (Elsevier/ScienceDirect)
            if 'sciencedirect' in pdf_url.lower() or 'elsevier' in pdf_url.lower():
                pdf_patterns.extend([
                    pdf_url.replace('/article/', '/article/pii/') + '/pdfft',
                    pdf_url.replace('/article/', '/article/pii/') + '/pdf',
                    pdf_url.replace('/article/', '/article/abstract/') + '/pdf',
                ])
            
            # Publisher-specific patterns (Springer)
            if 'springer' in pdf_url.lower():
                pdf_patterns.extend([
                    pdf_url.replace('/article/', '/article/pdf/') + '.pdf',
                    pdf_url.replace('/chapter/', '/chapter/pdf/'),
                    pdf_url + '.pdf',
                ])
            
            # Publisher-specific patterns (IEEE)
            if 'ieee' in pdf_url.lower():
                pdf_patterns.extend([
                    pdf_url.replace('/view/', '/download/') + '/pdf',
                    pdf_url.replace('/view/', '/download/') + '/file',
                ])
            
            # ArXiv patterns
            if 'arxiv.org' in pdf_url.lower():
                pdf_patterns.extend([
                    pdf_url.replace('/abs/', '/pdf/') + '.pdf',
                    pdf_url.replace('/abs/', '/pdf/'),
                ])
            
            # Generic patterns
            pdf_patterns.extend([
                pdf_url + '/pdf',
                pdf_url.replace('.html', '.pdf'),
                pdf_url.replace('.htm', '.pdf'),
                pdf_url + '.pdf',
                pdf_url.replace('/full/', '/pdf/'),
                pdf_url.replace('/abstract/', '/pdf/'),
                pdf_url.replace('/summary/', '/pdf/'),
            ])
            
            # Remove duplicates and try each pattern
            seen = set()
            for pattern in pdf_patterns:
                if pattern not in seen and pattern != pdf_url:
                    seen.add(pattern)
                    if self._try_download_pdf(pattern, filename, verbose):
                        return True
        
        # Try to extract PDF from HTML page (for DOI landing pages)
        try:
            response = self.session.get(pdf_url, timeout=30, allow_redirects=True)
            if response.status_code == 200:
                content_type = response.headers.get('content-type', '').lower()
                if 'html' in content_type:
                    # Try to extract PDF link from HTML
                    html_content = response.text
                    extracted_pdf = self._extract_pdf_from_html(html_content, pdf_url)
                    if extracted_pdf:
                        if self._try_download_pdf(extracted_pdf, filename, verbose):
                            return True
                    
                    # Also try common PDF button/link patterns in HTML
                    import re
                    # Look for PDF download links in common formats
                    pdf_link_patterns = [
                        r'href=["\']([^"\']*download[^"\']*\.pdf[^"\']*)["\']',
                        r'href=["\']([^"\']*\.pdf[^"\']*download[^"\']*)["\']',
                        r'href=["\']([^"\']*pdf[^"\']*)["\']',
                        r'data-pdf=["\']([^"\']*)["\']',
                        r'pdf-url=["\']([^"\']*)["\']',
                        r'data-url=["\']([^"\']*pdf[^"\']*)["\']',
                        r'["\']([^"\']*\.pdf)["\']',
                    ]
                    
                    for pattern in pdf_link_patterns:
                        matches = re.findall(pattern, html_content, re.IGNORECASE)
                        for match in matches:
                            if match and ('pdf' in match.lower() or match.endswith('.pdf')):
                                from urllib.parse import urljoin
                                pdf_link = urljoin(pdf_url, match)
                                if self._try_download_pdf(pdf_link, filename, verbose):
                                    return True
        except Exception:
            pass
        
        return False
    
    def _try_download_pdf(self, url: str, filename: str, verbose: bool = False) -> bool:
        """
        Try to download a PDF from a specific URL
        
        Args:
            url: URL to try
            filename: Local filename to save the PDF
            verbose: If True, log errors for debugging
            
        Returns:
            True if successful, False otherwise
        """
        try:
            response = self.session.get(url, timeout=30, stream=True, allow_redirects=True)
            
            # Check for HTTP errors
            if response.status_code == 403:
                # Forbidden - likely behind paywall or requires authentication
                if verbose:
                    print(f"    [403 Forbidden] {url[:80]}...")
                return False
            elif response.status_code == 404:
                # Not found
                if verbose:
                    print(f"    [404 Not Found] {url[:80]}...")
                return False
            elif response.status_code >= 400:
                # Other HTTP errors
                if verbose:
                    print(f"    [HTTP {response.status_code}] {url[:80]}...")
                return False
            
            # Check content type
            content_type = response.headers.get('content-type', '').lower()
            
            # Download content
            content = b''
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    content += chunk
                    # Check first bytes early if we have enough
                    if len(content) >= 4:
                        if content[:4] != b'%PDF':
                            # Not a PDF
                            if 'pdf' not in content_type and 'application/pdf' not in content_type:
                                # Check if it's HTML (likely a landing page)
                                if 'html' in content_type or content[:5] == b'<!DOC' or content[:5] == b'<html':
                                    # It's HTML, not PDF - this is expected for landing pages
                                    return False
                                # Might be a different file type
                                return False
            
            # Final check - must start with %PDF
            if len(content) < 4 or content[:4] != b'%PDF':
                if 'pdf' not in content_type and 'application/pdf' not in content_type:
                    # Check if it's HTML
                    if len(content) > 0 and (content[:5] == b'<!DOC' or content[:5] == b'<html' or b'<html' in content[:100].lower()):
                        return False
                    return False
            
            # Save the file
            with open(filename, 'wb') as f:
                f.write(content)
            
            return True
        except requests.exceptions.Timeout:
            if verbose:
                print(f"    [Timeout] {url[:80]}...")
            return False
        except requests.exceptions.RequestException as e:
            if verbose:
                print(f"    [Request Error] {url[:80]}... - {str(e)[:50]}")
            return False
        except Exception as e:
            if verbose:
                print(f"    [Error] {url[:80]}... - {str(e)[:50]}")
            return False
    
    def get_pdf_url_from_work(self, work: Dict) -> Optional[str]:
        """
        Extract PDF URL from OpenAlex work data
        Tries multiple methods to find PDF URLs, prioritizing open access sources
        
        Args:
            work: Work dictionary from OpenAlex
            
        Returns:
            PDF URL if available, None otherwise
        """
        # Method 1: Check best_oa_location first (best open access location)
        best_oa = work.get('best_oa_location')
        if best_oa:
            if best_oa.get('pdf_url'):
                return best_oa['pdf_url']
            # Also try landing_page_url from best_oa_location
            if best_oa.get('landing_page_url'):
                return best_oa['landing_page_url']
        
        # Method 2: Check all locations array (prioritize those with pdf_url)
        locations = work.get('locations', [])
        # First, try locations with direct PDF URLs
        for location in locations:
            if location.get('pdf_url'):
                return location['pdf_url']
        
        # Method 3: Check primary_location
        primary_location = work.get('primary_location')
        if primary_location:
            if primary_location.get('pdf_url'):
                return primary_location['pdf_url']
            # Also try landing_page_url from primary_location
            if primary_location.get('landing_page_url'):
                return primary_location['landing_page_url']
        
        # Method 4: For open access papers, try oa_url and DOI
        if work.get('open_access', {}).get('is_oa'):
            oa_url = work.get('open_access', {}).get('oa_url')
            if oa_url:
                return oa_url
            
            # Try DOI-based patterns
            doi = work.get('doi')
            if doi:
                return doi
        
        # Method 5: Try landing_page_url from locations array (if not already tried)
        for location in locations:
            if location.get('landing_page_url'):
                return location['landing_page_url']
        
        # Method 6: Try landing_page_url from best_oa_location (if not already tried)
        if best_oa and best_oa.get('landing_page_url'):
            return best_oa['landing_page_url']
        
        # Method 7: Try DOI as last resort
        doi = work.get('doi')
        if doi:
            return doi
        
        return None
    
    def get_all_pdf_urls_from_work(self, work: Dict) -> List[str]:
        """
        Extract ALL possible PDF URLs from OpenAlex work data
        Returns a list of all URLs that could potentially lead to a PDF
        
        Args:
            work: Work dictionary from OpenAlex
            
        Returns:
            List of all possible PDF/landing page URLs, prioritized by likelihood of success
        """
        urls = []
        seen = set()
        
        def add_url(url):
            """Helper to add URL if not already seen"""
            if url and url not in seen:
                seen.add(url)
                urls.append(url)
        
        # Priority 1: Direct PDF URLs from best_oa_location (most likely to work)
        best_oa = work.get('best_oa_location')
        if best_oa:
            if best_oa.get('pdf_url'):
                add_url(best_oa['pdf_url'])
            if best_oa.get('landing_page_url'):
                add_url(best_oa['landing_page_url'])
        
        # Priority 2: Direct PDF URLs from all locations
        locations = work.get('locations', [])
        for location in locations:
            if location.get('pdf_url'):
                add_url(location['pdf_url'])
        
        # Priority 3: Direct PDF URL from primary_location
        primary_location = work.get('primary_location')
        if primary_location:
            if primary_location.get('pdf_url'):
                add_url(primary_location['pdf_url'])
            if primary_location.get('landing_page_url'):
                add_url(primary_location['landing_page_url'])
        
        # Priority 4: Open access URL
        if work.get('open_access', {}).get('is_oa'):
            oa_url = work.get('open_access', {}).get('oa_url')
            if oa_url:
                add_url(oa_url)
        
        # Priority 5: Landing page URLs from all locations
        for location in locations:
            if location.get('landing_page_url'):
                add_url(location['landing_page_url'])
        
        # Priority 6: DOI (can be used to construct URLs)
        doi = work.get('doi')
        if doi:
            # Add DOI as-is
            add_url(doi)
            # Also add DOI with https://doi.org/ prefix if not already present
            if not doi.startswith('http'):
                add_url(f"https://doi.org/{doi}")
        
        return urls
    
    def download_work_pdfs(self, works: List[Dict], output_dir: str) -> int:
        """
        Download actual PDF documents for works
        Tries ALL available URLs for each work until one succeeds
        
        Args:
            works: List of work dictionaries from OpenAlex
            output_dir: Directory to save PDF files
            
        Returns:
            Number of PDFs successfully downloaded
        """
        import os
        import re
        
        os.makedirs(output_dir, exist_ok=True)
        downloaded = 0
        
        for idx, work in enumerate(works, 1):
            # Get ALL possible URLs for this work
            all_urls = self.get_all_pdf_urls_from_work(work)
            
            if not all_urls:
                # No URLs available at all
                title = work.get('title', 'Unknown')[:50]
                print(f"  ✗ No URLs available: {title}...")
                continue
            
            # Create safe filename from title or DOI
            title = work.get('title', f'document_{idx}')
            # Clean filename
            safe_title = re.sub(r'[<>:"/\\|?*]', '', title)
            safe_title = safe_title[:100]  # Limit length
            
            # Try to get DOI for filename
            doi = work.get('doi', '')
            if doi:
                doi_clean = doi.replace('https://doi.org/', '').replace('http://dx.doi.org/', '').replace('/', '_')
                filename = f"{idx:03d}_{doi_clean}.pdf"
            else:
                filename = f"{idx:03d}_{safe_title[:50]}.pdf"
            
            filepath = os.path.join(output_dir, filename)
            
            # Check if work is open access (helps prioritize)
            is_oa = work.get('open_access', {}).get('is_oa', False)
            oa_status = work.get('open_access', {}).get('oa_status', 'unknown')
            
            # Try each URL until one succeeds
            success = False
            verbose = idx > len(works) - 5  # Verbose for last 5 items
            
            for url_idx, pdf_url in enumerate(all_urls, 1):
                if self.download_pdf(pdf_url, filepath, verbose=verbose):
                    downloaded += 1
                    success = True
                    print(f"  ✓ Downloaded: {filename} (from URL {url_idx}/{len(all_urls)})")
                    break
                # Small delay between URL attempts
                time.sleep(0.1)
            
            if not success:
                # Remove failed download if it exists
                if os.path.exists(filepath):
                    os.remove(filepath)
                # Print which work failed (for debugging)
                title = work.get('title', 'Unknown')[:50]
                oa_info = f" [OA: {oa_status}]" if is_oa else " [Not OA]"
                url_count = len(all_urls)
                print(f"  ✗ Failed to download: {title}...{oa_info} (tried {url_count} URL(s))")
            
            # Small delay between downloads to be respectful
            time.sleep(0.2)
        
        return downloaded
    
    def save_to_markdown(self, data: List[Dict], filename: str):
        """
        Save scraped data to a Markdown file (.md)
        
        Args:
            data: List of dictionaries to save
            filename: Output filename (should end with .md)
        """
        if not data:
            print("No data to save")
            return
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write('# OpenAlex Scraped Data\n\n')
                f.write(f'**Total Records:** {len(data)}\n\n')
                f.write('---\n\n')
                
                for idx, item in enumerate(data, 1):
                    f.write(f'## Record {idx}\n\n')
                    
                    for key, value in item.items():
                        if value is None:
                            continue
                        
                        # Format the value
                        if isinstance(value, (dict, list)):
                            value_str = '```json\n' + json.dumps(value, indent=2, ensure_ascii=False) + '\n```'
                        else:
                            value_str = str(value)
                        
                        f.write(f'**{key}:**\n\n')
                        f.write(f'{value_str}\n\n')
                    
                    f.write('---\n\n')
            
            print(f"Saved {len(data)} records to {filename}")
        except Exception as e:
            print(f"Error saving to Markdown: {e}")
    
    def save_to_html(self, data: List[Dict], filename: str):
        """
        Save scraped data to an HTML file
        
        Args:
            data: List of dictionaries to save
            filename: Output filename (should end with .html)
        """
        if not data:
            print("No data to save")
            return
        
        try:
            html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>OpenAlex Scraped Data</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            margin: 20px;
            background-color: #f5f5f5;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background-color: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        h1 {
            color: #333;
            border-bottom: 3px solid #4CAF50;
            padding-bottom: 10px;
        }
        .record {
            margin: 20px 0;
            padding: 15px;
            border: 1px solid #ddd;
            border-radius: 5px;
            background-color: #fafafa;
        }
        .record h2 {
            color: #4CAF50;
            margin-top: 0;
        }
        .field {
            margin: 10px 0;
        }
        .field-label {
            font-weight: bold;
            color: #555;
        }
        .field-value {
            margin-left: 10px;
            color: #333;
            word-wrap: break-word;
        }
        pre {
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 4px;
            overflow-x: auto;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>OpenAlex Scraped Data</h1>
        <p><strong>Total Records:</strong> {total_records}</p>
"""
            
            for idx, item in enumerate(data, 1):
                html_content += '        <div class="record">\n'
                html_content += f'            <h2>Record {idx}</h2>\n'
                
                for key, value in item.items():
                    if value is None:
                        continue
                    
                    # Format the value
                    if isinstance(value, (dict, list)):
                        value_str = json.dumps(value, indent=2, ensure_ascii=False)
                        value_str = value_str.replace('<', '&lt;').replace('>', '&gt;')
                        value_html = f'<pre>{value_str}</pre>'
                    else:
                        value_str = str(value).replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                        value_html = value_str
                    
                    key_escaped = str(key).replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                    html_content += '            <div class="field">\n'
                    html_content += f'                <span class="field-label">{key_escaped}:</span>\n'
                    html_content += f'                <span class="field-value">{value_html}</span>\n'
                    html_content += '            </div>\n'
                
                html_content += '        </div>\n'
            
            html_content += """    </div>
</body>
</html>"""
            
            html_content = html_content.format(total_records=len(data))
            
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            print(f"Saved {len(data)} records to {filename}")
        except Exception as e:
            print(f"Error saving to HTML: {e}")


def scrape_keywords_from_csv(csv_filename: str, 
                              max_results_per_keyword: int = 50,
                              output_dir: str = "output",
                              save_combined: bool = True,
                              organize_by_keyword: bool = True):
    """
    Scrape OpenAlex data for keywords from a CSV file
    For each keyword, finds all related documents/works from OpenAlex and saves them in all formats
    
    Args:
        csv_filename: Path to CSV file containing keywords (one per line)
        max_results_per_keyword: Maximum number of results to fetch per keyword (None for all)
        output_dir: Directory to save output files
        save_combined: Whether to save all results in a single combined file
        organize_by_keyword: If True, creates a subfolder for each keyword
    
    Returns:
        Dictionary mapping keywords to their scraped works
    """
    import csv
    import os
    
    # Initialize scraper
    scraper = OpenAlexScraper(email="vipuldholariya1991dev@gmail.com")
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Read keywords from CSV or text file
    keywords = []
    try:
        with open(csv_filename, 'r', encoding='utf-8') as f:
            # Try reading as CSV first
            content = f.read()
            f.seek(0)  # Reset file pointer
            
            # Check if it's a CSV with commas or just plain text (one per line)
            if ',' in content:
                # CSV format
                reader = csv.reader(f)
                for row in reader:
                    if row and row[0].strip():  # Skip empty rows
                        keywords.append(row[0].strip())
            else:
                # Plain text format (one keyword per line)
                for line in f:
                    keyword = line.strip()
                    if keyword:  # Skip empty lines
                        keywords.append(keyword)
    except FileNotFoundError:
        print(f"Error: File '{csv_filename}' not found.")
        return {}
    except Exception as e:
        print(f"Error reading file: {e}")
        return {}
    
    if not keywords:
        print("No keywords found in CSV file.")
        return {}
    
    print(f"=== Scraping OpenAlex for {len(keywords)} keywords ===\n")
    print(f"For each keyword, finding all related documents/works from OpenAlex...\n")
    
    all_results = {}
    all_works_combined = []
    keyword_summaries = {}
    
    # Scrape data for each keyword
    for idx, keyword in enumerate(keywords, 1):
        print(f"\n{'='*60}")
        print(f"[{idx}/{len(keywords)}] Keyword: '{keyword}'")
        print(f"{'='*60}")
        
        try:
            # Enhance search query for boiler-related keywords
            # Add "boiler" to search if keyword doesn't already contain it
            search_query = keyword
            keyword_lower = keyword.lower()
            
            # Check if keyword is boiler-related
            boiler_terms = ['boiler', 'whrb', 'waste heat', 'heat exchanger', 'heat recovery']
            is_boiler_related = any(term in keyword_lower for term in boiler_terms)
            
            # If keyword doesn't contain boiler terms but is about overheating/trip problems,
            # add boiler context to search
            if not is_boiler_related and ('overheating' in keyword_lower or 'trip' in keyword_lower or 'startup' in keyword_lower):
                search_query = f"{keyword} boiler"
                print(f"Enhanced search query: '{search_query}' (adding boiler context)")
            
            # Search for all works/documents related to this keyword (English only)
            print(f"Searching OpenAlex for documents related to '{search_query}' (English only)...")
            # Filter for English language only - OpenAlex uses 'language:en' format
            works = scraper.get_all_works(
                query=search_query,
                filter_params={'language': 'en'},  # Filter for English language only
                max_results=max_results_per_keyword,
                use_cursor=True
            )
            
            # Filter results to ensure they're boiler-related
            if works and not is_boiler_related:
                # Check if results contain boiler-related terms in title or abstract
                boiler_keywords = ['boiler', 'heat exchanger', 'waste heat', 'whrb', 'steam generator', 
                                 'heat recovery', 'hrsg', 'heat transfer', 'thermal']
                filtered_works = []
                for work in works:
                    title = work.get('title', '').lower()
                    # Check abstract if available
                    abstract = ''
                    if work.get('abstract_inverted_index'):
                        # Reconstruct abstract from inverted index (simplified)
                        abstract = ' '.join(work.get('abstract_inverted_index', {}).keys()).lower()
                    
                    # Check if work is boiler-related
                    work_text = f"{title} {abstract}"
                    if any(term in work_text for term in boiler_keywords):
                        filtered_works.append(work)
                
                if len(filtered_works) < len(works):
                    print(f"  Filtered {len(works)} results to {len(filtered_works)} boiler-related documents")
                    works = filtered_works
            
            print(f"✓ Found {len(works)} documents/works for '{keyword}'")
            
            # Add keyword to each work for tracking
            for work in works:
                work['search_keyword'] = keyword
            
            all_results[keyword] = works
            all_works_combined.extend(works)
            
            # Create safe filename from keyword
            keyword_safe = "".join(c for c in keyword if c.isalnum() or c in (' ', '-', '_')).strip()
            keyword_safe = keyword_safe.replace(' ', '_')
            
            # Determine output path
            if organize_by_keyword:
                keyword_dir = os.path.join(output_dir, keyword_safe)
                os.makedirs(keyword_dir, exist_ok=True)
                base_filename = os.path.join(keyword_dir, "documents")
            else:
                base_filename = os.path.join(output_dir, f"{keyword_safe}_documents")
            
            if works:
                print(f"\nSaving {len(works)} documents for '{keyword}' in CSV format...")
                
                # Save CSV
                scraper.save_to_csv(works, f"{base_filename}.csv")
                
                # Download actual PDF documents
                print(f"\nDownloading PDF documents for '{keyword}'...")
                pdf_dir = os.path.join(keyword_dir if organize_by_keyword else output_dir, f"{keyword_safe}_pdfs")
                downloaded_count = scraper.download_work_pdfs(works, pdf_dir)
                print(f"  Downloaded {downloaded_count} out of {len(works)} PDF documents")
                
                # Store summary
                keyword_summaries[keyword] = {
                    'count': len(works),
                    'pdf_count': downloaded_count,
                    'path': keyword_dir if organize_by_keyword else os.path.dirname(base_filename),
                    'csv_file': f"{base_filename}.csv",
                    'pdf_dir': pdf_dir
                }
                
                print(f"\n✓ All documents saved for '{keyword}':")
                if organize_by_keyword:
                    print(f"   Location: {keyword_dir}/")
                    print(f"   - CSV: {base_filename}.csv")
                    print(f"   - PDFs: {pdf_dir}/ ({downloaded_count} files)")
                else:
                    print(f"   Location: {os.path.dirname(base_filename)}/")
                    print(f"   - CSV: {base_filename}.csv")
                    print(f"   - PDFs: {pdf_dir}/ ({downloaded_count} files)")
            else:
                print(f"⚠ No documents found for '{keyword}'")
                keyword_summaries[keyword] = {'count': 0, 'path': None, 'files': []}
            
        except Exception as e:
            print(f"   ✗ Error scraping '{keyword}': {e}\n")
            all_results[keyword] = []
            keyword_summaries[keyword] = {'count': 0, 'path': None, 'files': []}
    
    # Save combined results in CSV format only
    if save_combined and all_works_combined:
        base_combined = os.path.join(output_dir, "all_keywords_combined")
        
        # Save CSV
        scraper.save_to_csv(all_works_combined, f"{base_combined}.csv")
        
        # Download all PDFs to combined folder
        print(f"\nDownloading all PDF documents to combined folder...")
        combined_pdf_dir = os.path.join(output_dir, "all_keywords_combined_pdfs")
        combined_downloaded = scraper.download_work_pdfs(all_works_combined, combined_pdf_dir)
        
        print(f"\n✓ Combined results saved:")
        print(f"   - {base_combined}.csv")
        print(f"   - PDFs: {combined_pdf_dir}/ ({combined_downloaded} files)")
        print(f"   Total works: {len(all_works_combined)}")
    
    # Print detailed summary
    print(f"\n\n{'='*60}")
    print(f"=== SCRAPING COMPLETE ===")
    print(f"{'='*60}\n")
    
    print(f"Keywords processed: {len(keywords)}")
    print(f"Total documents found: {len(all_works_combined)}\n")
    
    print("Summary by keyword:")
    print("-" * 60)
    for keyword, summary in keyword_summaries.items():
        print(f"  • {keyword}: {summary['count']} documents")
        if summary['path']:
            print(f"    Saved in: {summary['path']}/")
    
    print(f"\nAll results saved in: {output_dir}/")
    
    if save_combined and all_works_combined:
        print(f"\nCombined file location: {output_dir}/all_keywords_combined.*")
    
    print(f"\n{'='*60}\n")
    
    return all_results


def main():
    """
    Main function - Scrapes OpenAlex data for keywords from keywords1.csv
    
    For each keyword in the CSV file:
    1. Searches OpenAlex for all related documents/works
    2. Saves metadata in CSV format
    3. Downloads actual PDF documents (where available) to separate folders
    4. Organizes documents by keyword in separate folders
    """
    # Scrape keywords from CSV file
    # For each of the 4 keywords, finds all related documents from OpenAlex
    scrape_keywords_from_csv(
        csv_filename="keywords1.csv",
        max_results_per_keyword=50,  # Set to None to get all results, or adjust number
        output_dir="output",
        save_combined=True,  # Also saves a combined file with all keywords
        organize_by_keyword=True  # Creates a folder for each keyword
    )


def main_demo():
    """Example usage of the OpenAlex scraper following official best practices"""
    
    # Initialize scraper
    # NO LOGIN OR API KEY REQUIRED - API is completely free!
    # Email is recommended for better response times (polite pool)
    # Optional: Add api_key parameter if you have Premium plan
    scraper = OpenAlexScraper(email="vipuldholariya1991dev@gmail.com")
    
    print("=== OpenAlex Scraper Demo ===\n")
    
    # Example 1: Search for works about "machine learning" with cursor pagination
    print("1. Searching for works about 'machine learning' (using cursor pagination)...")
    works = scraper.get_all_works(
        query="machine learning",
        max_results=10,
        use_cursor=True  # More efficient than page-based
    )
    print(f"   Found {len(works)} works")
    
    if works:
        print(f"\n   First work:")
        work = works[0]
        print(f"   Title: {work.get('title', 'N/A')}")
        print(f"   DOI: {work.get('doi', 'N/A')}")
        print(f"   Publication Year: {work.get('publication_year', 'N/A')}")
        print(f"   Citations: {work.get('cited_by_count', 0)}")
    
    # Example 2: Get a specific work
    print("\n2. Fetching a specific work...")
    work = scraper.get_work("W2741803907")
    if work:
        print(f"   Title: {work.get('title', 'N/A')}")
        print(f"   Authors: {len(work.get('authorships', []))} author(s)")
    
    # Example 3: Search for authors with filtering
    print("\n3. Searching for authors named 'Einstein'...")
    authors_result = scraper.search_authors(query="Einstein", per_page=5)
    if 'results' in authors_result:
        print(f"   Found {len(authors_result['results'])} authors")
        for author in authors_result['results'][:3]:
            print(f"   - {author.get('display_name', 'N/A')}")
    
    # Example 4: Get works by a specific author using filters
    print("\n4. Getting works by a specific author (using filters)...")
    author_works = scraper.get_works_by_author("A2208157607", max_results=5)
    print(f"   Found {len(author_works)} works")
    
    # Example 5: Search with sorting (most cited first)
    print("\n5. Searching for works sorted by citations (descending)...")
    sorted_works = scraper.search_works(
        query="deep learning",
        sort_params={'cited_by_count': 'desc'},
        per_page=5
    )
    if 'results' in sorted_works:
        print(f"   Found {len(sorted_works['results'])} works")
        for work in sorted_works['results'][:3]:
            print(f"   - {work.get('title', 'N/A')[:60]}... ({work.get('cited_by_count', 0)} citations)")
    
    # Example 6: Get works from a specific year
    print("\n6. Getting works from 2023...")
    works_2023 = scraper.get_works_by_year(2023, max_results=5)
    print(f"   Found {len(works_2023)} works")
    
    # Example 7: Search for institutions
    print("\n7. Searching for institutions...")
    institutions = scraper.search_institutions(query="MIT", per_page=3)
    if 'results' in institutions:
        print(f"   Found {len(institutions['results'])} institutions")
        for inst in institutions['results']:
            print(f"   - {inst.get('display_name', 'N/A')}")
    
    # Example 8: Get a topic
    print("\n8. Getting a specific topic...")
    topic = scraper.get_topic("T11536")
    if topic:
        print(f"   Topic: {topic.get('display_name', 'N/A')}")
        print(f"   Works count: {topic.get('works_count', 0)}")
    
    # Save results to files
    if works:
        scraper.save_to_json(works, "scraped_works.json")
        scraper.save_to_csv(works, "scraped_works.csv")
        print("\n✓ Results saved to scraped_works.json and scraped_works.csv")
    
    print("\n=== Demo Complete ===")
    print("\nTips:")
    print("- NO LOGIN OR API KEY REQUIRED - API is completely free!")
    print("- Include your email for better response times (polite pool)")
    print("- Use cursor-based pagination for better performance")
    print("- Use filters to narrow down results efficiently")
    print("- Sort results to get the most relevant data first")
    print("- Optional: Premium plan with API key for higher limits and special features")


if __name__ == "__main__":
    main()

